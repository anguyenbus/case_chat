"""
Document parser for PDF, TXT, and DOCX files.

Extracts text and metadata from supported document formats.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Final

logger = logging.getLogger(__name__)


# ============================================================================
# Document Parser
# ============================================================================


class DocumentParser:
    """
    Parse documents and extract text and metadata.

    Supports PDF, TXT, and DOCX file formats.
    Extracts text content and basic metadata (filename, type, size).

    Attributes
    ----------
    _allowed_types : set[str]
        Set of allowed file extensions

    """

    __slots__ = ("_allowed_types",)

    ALLOWED_TYPES: Final[set[str]] = {".pdf", ".txt", ".docx"}

    def __init__(self):
        """Initialize the DocumentParser."""
        self._allowed_types = self.ALLOWED_TYPES

    def detect_file_type(self, filename: str) -> str:
        """
        Detect file type from filename extension.

        Parameters
        ----------
        filename : str
            Name of the file

        Returns
        -------
        str
            File extension (e.g., '.pdf', '.txt', '.docx')

        Raises
        ------
        ValueError
            If file type is not supported

        Examples
        --------
        >>> parser = DocumentParser()
        >>> parser.detect_file_type("document.pdf")
        '.pdf'

        """
        ext = Path(filename).suffix.lower()

        if ext not in self._allowed_types:
            raise ValueError(
                f"Unsupported file type: {ext}. Allowed types: {', '.join(self._allowed_types)}"
            )

        return ext

    def parse_txt_string(
        self,
        content: str,
        filename: str,
    ) -> tuple[str, dict[str, any]]:
        """
        Parse TXT content from string.

        Parameters
        ----------
        content : str
            Text content
        filename : str
            Original filename

        Returns
        -------
        tuple[str, dict[str, any]]
            Tuple of (text_content, metadata)

        Examples
        --------
        >>> parser = DocumentParser()
        >>> text, meta = parser.parse_txt_string("Hello world", "test.txt")
        >>> text
        'Hello world'

        """
        # Encode to get file size
        file_size = len(content.encode("utf-8"))

        # Create metadata
        metadata = {
            "filename": filename,
            "file_type": ".txt",
            "file_size": file_size,
        }

        logger.debug(f"Parsed TXT file {filename} ({file_size} bytes)")

        return content, metadata

    def parse_txt(self, file_path: Path) -> tuple[str, dict[str, any]]:
        """
        Parse TXT file and extract text.

        Parameters
        ----------
        file_path : Path
            Path to TXT file

        Returns
        -------
        tuple[str, dict[str, any]]
            Tuple of (text_content, metadata)

        Raises
        ------
        FileNotFoundError
            If file does not exist
        ValueError
            If file type is not TXT

        Examples
        --------
        >>> parser = DocumentParser()
        >>> text, meta = parser.parse_txt(Path("document.txt"))

        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file type
        file_type = self.detect_file_type(file_path.name)
        if file_type != ".txt":
            raise ValueError(f"Expected .txt file, got {file_type}")

        # Read file with UTF-8 encoding
        try:
            with open(file_path, encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1
            with open(file_path, encoding="latin-1") as f:
                text = f.read()

        # Get file size
        file_size = file_path.stat().st_size

        # Create metadata
        metadata = {
            "filename": file_path.name,
            "file_type": ".txt",
            "file_size": file_size,
        }

        logger.debug(f"Parsed TXT file {file_path.name} ({file_size} bytes)")

        return text, metadata

    def parse_pdf(self, file_path: Path) -> tuple[str, dict[str, any]]:
        """
        Parse PDF file and extract text.

        Parameters
        ----------
        file_path : Path
            Path to PDF file

        Returns
        -------
        tuple[str, dict[str, any]]
            Tuple of (text_content, metadata)

        Raises
        ------
        FileNotFoundError
            If file does not exist
        ValueError
            If file type is not PDF or parsing fails

        Examples
        --------
        >>> parser = DocumentParser()
        >>> text, meta = parser.parse_pdf(Path("document.pdf"))

        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file type
        file_type = self.detect_file_type(file_path.name)
        if file_type != ".pdf":
            raise ValueError(f"Expected .pdf file, got {file_type}")

        try:
            # Import pypdf
            from pypdf import PdfReader

            # Create PDF reader
            reader = PdfReader(str(file_path))

            # Extract text from all pages
            text_parts = []
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page: {e}")
                    continue

            text = "\n\n".join(text_parts)

            # Get file size
            file_size = file_path.stat().st_size

            # Create metadata
            metadata = {
                "filename": file_path.name,
                "file_type": ".pdf",
                "file_size": file_size,
                "page_count": len(reader.pages),
            }

            logger.debug(
                f"Parsed PDF file {file_path.name} ({file_size} bytes, {len(reader.pages)} pages)"
            )

            return text, metadata

        except ImportError:
            raise ValueError("pypdf is not installed. Install with: pip install pypdf") from None
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file: {e}") from e

    def parse_docx(self, file_path: Path) -> tuple[str, dict[str, any]]:
        """
        Parse DOCX file and extract text.

        Parameters
        ----------
        file_path : Path
            Path to DOCX file

        Returns
        -------
        tuple[str, dict[str, any]]
            Tuple of (text_content, metadata)

        Raises
        ------
        FileNotFoundError
            If file does not exist
        ValueError
            If file type is not DOCX or parsing fails

        Examples
        --------
        >>> parser = DocumentParser()
        >>> text, meta = parser.parse_docx(Path("document.docx"))

        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Validate file type
        file_type = self.detect_file_type(file_path.name)
        if file_type != ".docx":
            raise ValueError(f"Expected .docx file, got {file_type}")

        try:
            # Import python-docx
            from docx import Document

            # Create document reader
            doc = Document(str(file_path))

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)

            text = "\n\n".join(text_parts)

            # Get file size
            file_size = file_path.stat().st_size

            # Create metadata
            metadata = {
                "filename": file_path.name,
                "file_type": ".docx",
                "file_size": file_size,
                "paragraph_count": len(doc.paragraphs),
            }

            logger.debug(
                f"Parsed DOCX file {file_path.name} "
                f"({file_size} bytes, {len(doc.paragraphs)} paragraphs)"
            )

            return text, metadata

        except ImportError:
            raise ValueError(
                "python-docx is not installed. Install with: pip install python-docx"
            ) from None
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX file: {e}") from e
